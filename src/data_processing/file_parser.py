"""
universal_parser.py — ULTRA PRO Edition

Uniwersalny parser plików z zaawansowanymi funkcjami:
- Wsparcie dla 15+ formatów (CSV/XLSX/JSON/Parquet/PDF/DOCX/etc.)
- Inteligentna detekcja encoding i delimitera
- Robust error handling z graceful degradation
- Memory-safe z limitami
- PyArrow backend support
- Comprehensive logging
- Thread-safe operations
"""

from __future__ import annotations

import io
import csv
import json
import gzip
import logging
import warnings as py_warnings
from dataclasses import dataclass, field
from typing import Tuple, Optional, Any, Dict, List, Union
from pathlib import Path

import pandas as pd
import numpy as np

# ========================================================================================
# OPTIONAL DEPENDENCIES (graceful fallback)
# ========================================================================================

try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    chardet = None  # type: ignore
    HAS_CHARDET = False

try:
    from PyPDF2 import PdfReader
    HAS_PDF = True
except ImportError:
    PdfReader = None  # type: ignore
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    Document = None  # type: ignore
    HAS_DOCX = False

try:
    import pyarrow.parquet as pq
    HAS_PYARROW = True
except ImportError:
    HAS_PYARROW = False

# ========================================================================================
# LOGGING
# ========================================================================================

def get_logger(name: str = "universal_parser") -> logging.Logger:
    """Konfiguruje logger."""
    logger = logging.getLogger(name)
    if not logger.handlers:
        logger.setLevel(logging.INFO)
        handler = logging.StreamHandler()
        handler.setFormatter(
            logging.Formatter(
                "%(asctime)s | %(levelname)s | %(name)s | %(message)s",
                datefmt="%Y-%m-%d %H:%M:%S"
            )
        )
        logger.addHandler(handler)
        logger.propagate = False
    return logger

LOGGER = get_logger()

# ========================================================================================
# CONSTANTS
# ========================================================================================

# Obsługiwane rozszerzenia
SUPPORTED_EXTENSIONS = {
    "csv", "tsv", "txt", "xlsx", "xls", "json", "jsonl", "ndjson",
    "docx", "pdf", "doc", "parquet", "feather", 
    "csv.gz", "tsv.gz", "json.gz"
}

# Domyślne limity bezpieczeństwa
DEFAULT_MAX_FILE_SIZE_MB = 500
DEFAULT_MAX_ROWS = 10_000_000
DEFAULT_MAX_COLS = 1000

# CSV detekcja
DEFAULT_CSV_DELIMITERS = (",", ";", "\t", "|", ":")
DEFAULT_CSV_ENCODINGS = ("utf-8", "utf-8-sig", "cp1250", "iso-8859-2", "latin-1", "windows-1252")
DEFAULT_SNIFF_BYTES = 10_000

# Quotingi do testowania
CSV_QUOTING_OPTIONS = (csv.QUOTE_MINIMAL, csv.QUOTE_ALL, csv.QUOTE_NONNUMERIC, csv.QUOTE_NONE)

# ========================================================================================
# DATACLASSES
# ========================================================================================

@dataclass(frozen=True)
class ParseOptions:
    """Opcje parsowania z rozszerzonymi możliwościami."""
    
    # PyArrow backend
    prefer_pyarrow: bool = True
    
    # Limity bezpieczeństwa
    max_rows: Optional[int] = None
    max_cols: Optional[int] = None
    max_file_size_mb: float = DEFAULT_MAX_FILE_SIZE_MB
    
    # Excel specyficzne
    excel_sheet: Optional[Union[str, int]] = None  # Nazwa lub indeks (0-based)
    excel_read_all_sheets: bool = False
    
    # JSON specyficzne
    json_force_lines: bool = False  # Wymuś lines=True
    json_normalize_nested: bool = True
    json_max_level: int = 1
    
    # CSV heurystyki
    csv_delimiters: Tuple[str, ...] = DEFAULT_CSV_DELIMITERS
    csv_try_encodings: Tuple[str, ...] = DEFAULT_CSV_ENCODINGS
    csv_sniff_bytes: int = DEFAULT_SNIFF_BYTES
    csv_skip_blank_lines: bool = True
    csv_on_bad_lines: str = "warn"  # "error", "warn", "skip"
    
    # Post-processing
    sanitize_inf: bool = True
    sanitize_nan_strings: bool = True  # "NaN", "nan", "null" -> pd.NA
    drop_empty_rows: bool = True
    drop_empty_cols: bool = False
    strip_whitespace: bool = True
    reset_index: bool = True
    
    # PDF/DOCX
    pdf_extract_images: bool = False
    docx_preserve_formatting: bool = False
    
    # Debug
    verbose: bool = False
    raise_on_warning: bool = False


@dataclass
class ParseResult:
    """Wynik parsowania z metadanymi."""
    
    df: Optional[pd.DataFrame] = None
    text: Optional[str] = None
    warnings: List[str] = field(default_factory=list)
    meta: Dict[str, Any] = field(default_factory=dict)
    
    @property
    def success(self) -> bool:
        """Czy parsowanie zakończyło się sukcesem."""
        return (self.df is not None and not self.df.empty) or bool(self.text)
    
    @property
    def is_tabular(self) -> bool:
        """Czy wynik to dane tabelaryczne."""
        return self.df is not None
    
    @property
    def is_text(self) -> bool:
        """Czy wynik to tekst."""
        return self.text is not None
    
    def add_warning(self, msg: str) -> None:
        """Dodaje ostrzeżenie."""
        self.warnings.append(msg)
        LOGGER.warning(msg)


# ========================================================================================
# VALIDATION & SAFETY
# ========================================================================================

def _validate_file_size(file_bytes: bytes, max_size_mb: float, filename: str) -> None:
    """
    Waliduje rozmiar pliku.
    
    Args:
        file_bytes: Dane binarne
        max_size_mb: Maksymalny rozmiar w MB
        filename: Nazwa pliku (do logowania)
        
    Raises:
        ValueError: Jeśli plik za duży
    """
    size_mb = len(file_bytes) / 1e6
    
    if size_mb > max_size_mb:
        raise ValueError(
            f"Plik '{filename}' przekracza limit rozmiaru: "
            f"{size_mb:.1f} MB > {max_size_mb:.1f} MB"
        )
    
    LOGGER.debug(f"File size OK: {size_mb:.2f} MB")


def _validate_dataframe(df: pd.DataFrame, opts: ParseOptions) -> pd.DataFrame:
    """
    Waliduje i sanityzuje DataFrame.
    
    Args:
        df: DataFrame do walidacji
        opts: Opcje parsowania
        
    Returns:
        Zwalidowany DataFrame
    """
    if df is None or not isinstance(df, pd.DataFrame):
        return df
    
    # Sprawdź czy pusty
    if df.empty:
        LOGGER.warning("DataFrame is empty")
        return df
    
    # Limity wierszy/kolumn
    original_shape = df.shape
    
    if opts.max_cols and df.shape[1] > opts.max_cols:
        LOGGER.warning(f"Limiting columns: {df.shape[1]} -> {opts.max_cols}")
        df = df.iloc[:, :opts.max_cols]
    
    if opts.max_rows and len(df) > opts.max_rows:
        LOGGER.warning(f"Limiting rows: {len(df)} -> {opts.max_rows}")
        df = df.head(opts.max_rows)
    
    # Post-processing
    if opts.drop_empty_rows:
        before = len(df)
        df = df.dropna(how="all")
        after = len(df)
        if before != after:
            LOGGER.debug(f"Dropped {before - after} empty rows")
    
    if opts.drop_empty_cols:
        before = df.shape[1]
        df = df.dropna(axis=1, how="all")
        after = df.shape[1]
        if before != after:
            LOGGER.debug(f"Dropped {before - after} empty columns")
    
    # Sanityzacja wartości
    if opts.sanitize_inf:
        # Zamień inf/-inf na NA
        df = df.replace([np.inf, -np.inf], pd.NA)
    
    if opts.sanitize_nan_strings:
        # Zamień stringowe reprezentacje NaN
        nan_strings = ["NaN", "nan", "null", "NULL", "None", "NONE", "#N/A", "N/A"]
        df = df.replace(nan_strings, pd.NA)
    
    if opts.strip_whitespace:
        # Strip whitespace z kolumn string
        for col in df.select_dtypes(include=["object", "string"]).columns:
            try:
                df[col] = df[col].str.strip()
            except Exception:
                pass
    
    if opts.reset_index:
        df = df.reset_index(drop=True)
    
    if df.shape != original_shape:
        LOGGER.debug(f"Shape changed: {original_shape} -> {df.shape}")
    
    return df


def _get_dtype_backend_kwargs(opts: ParseOptions) -> Dict[str, Any]:
    """
    Zwraca kwargs dla PyArrow backend.
    
    Args:
        opts: Opcje parsowania
        
    Returns:
        Dict z dtype_backend lub pusty dict
    """
    if not opts.prefer_pyarrow:
        return {}
    
    if not HAS_PYARROW:
        if opts.verbose:
            LOGGER.info("PyArrow not available, using default backend")
        return {}
    
    # Sprawdź wersję pandas
    try:
        pd_version = tuple(map(int, pd.__version__.split(".")[:2]))
        
        if pd_version >= (2, 0):
            return {"dtype_backend": "pyarrow"}
        else:
            if opts.verbose:
                LOGGER.info(f"PyArrow backend requires pandas>=2.0 (current: {pd.__version__})")
            return {}
    except Exception:
        return {}


# ========================================================================================
# ENCODING & DELIMITER DETECTION
# ========================================================================================

def _detect_encoding(
    data: bytes,
    candidates: Tuple[str, ...],
    sniff_bytes: int
) -> str:
    """
    Wykrywa encoding pliku.
    
    Args:
        data: Dane binarne
        candidates: Lista encoding do sprawdzenia
        sniff_bytes: Liczba bajtów do analizy
        
    Returns:
        Wykryty encoding
    """
    sample = data[:sniff_bytes]
    
    # Próba 1: chardet (jeśli dostępny)
    if HAS_CHARDET and chardet is not None:
        try:
            result = chardet.detect(sample)
            
            if result and result.get("encoding"):
                encoding = result["encoding"].lower()
                confidence = result.get("confidence", 0)
                
                if confidence > 0.7:
                    LOGGER.debug(f"Chardet detected: {encoding} (confidence: {confidence:.2%})")
                    return encoding
        except Exception as e:
            LOGGER.debug(f"Chardet failed: {e}")
    
    # Próba 2: BOM detection
    if sample.startswith(b"\xef\xbb\xbf"):
        LOGGER.debug("UTF-8 BOM detected")
        return "utf-8-sig"
    
    if sample.startswith(b"\xff\xfe") or sample.startswith(b"\xfe\xff"):
        LOGGER.debug("UTF-16 BOM detected")
        return "utf-16"
    
    # Próba 3: Sekwencyjna próba dekodowania
    for encoding in candidates:
        try:
            sample.decode(encoding)
            LOGGER.debug(f"Encoding detected: {encoding}")
            return encoding
        except (UnicodeDecodeError, LookupError):
            continue
    
    # Fallback
    LOGGER.warning("Could not detect encoding, using utf-8")
    return "utf-8"


def _sniff_delimiter(
    sample_text: str,
    delimiters: Tuple[str, ...]
) -> str:
    """
    Wykrywa delimiter w CSV.
    
    Args:
        sample_text: Próbka tekstu
        delimiters: Możliwe delimitery
        
    Returns:
        Wykryty delimiter
    """
    # Heurystyka 1: Dużo tabów = TSV
    tab_count = sample_text.count("\t")
    comma_count = sample_text.count(",")
    semicolon_count = sample_text.count(";")
    pipe_count = sample_text.count("|")
    
    counts = {
        "\t": tab_count,
        ",": comma_count,
        ";": semicolon_count,
        "|": pipe_count
    }
    
    # Jeśli tabulatory są dominujące
    if tab_count > max(comma_count, semicolon_count, pipe_count) and tab_count >= 2:
        LOGGER.debug("Delimiter detected (heuristic): TAB")
        return "\t"
    
    # Heurystyka 2: csv.Sniffer
    try:
        dialect = csv.Sniffer().sniff(sample_text, delimiters="".join(delimiters))
        detected = dialect.delimiter
        LOGGER.debug(f"Delimiter detected (Sniffer): {repr(detected)}")
        return detected
    except Exception as e:
        LOGGER.debug(f"CSV Sniffer failed: {e}")
    
    # Heurystyka 3: Najwięcej wystąpień
    if counts:
        most_common = max(counts.items(), key=lambda x: x[1])
        
        if most_common[1] >= 2:  # Co najmniej 2 wystąpienia
            LOGGER.debug(f"Delimiter detected (count): {repr(most_common[0])}")
            return most_common[0]
    
    # Fallback: przecinek
    LOGGER.warning("Could not detect delimiter, using comma")
    return ","


# ========================================================================================
# FORMAT-SPECIFIC READERS
# ========================================================================================

def _read_csv_robust(
    data: bytes,
    sep_hint: Optional[str],
    opts: ParseOptions
) -> pd.DataFrame:
    """
    Robust CSV reader z wieloma fallbackami.
    
    Args:
        data: Dane binarne
        sep_hint: Podpowiedź separatora (None = auto-detect)
        opts: Opcje parsowania
        
    Returns:
        DataFrame
        
    Raises:
        ValueError: Jeśli parsowanie się nie uda
    """
    # Detekcja encoding
    encoding = _detect_encoding(data, opts.csv_try_encodings, opts.csv_sniff_bytes)
    
    # Próbka dla sniffingu
    try:
        sample = data[:opts.csv_sniff_bytes].decode(encoding, errors="ignore")
    except Exception:
        sample = data[:opts.csv_sniff_bytes].decode("utf-8", errors="ignore")
    
    # Detekcja separatora
    sep = sep_hint or _sniff_delimiter(sample, opts.csv_delimiters)
    
    # Kwargs dla read_csv
    base_kwargs = {
        "sep": sep,
        "encoding": encoding,
        "low_memory": False,
        "skip_blank_lines": opts.csv_skip_blank_lines,
        "on_bad_lines": opts.csv_on_bad_lines,
        **_get_dtype_backend_kwargs(opts)
    }
    
    # Próba 1: Standardowe parsowanie z różnymi quotingami
    for quoting in CSV_QUOTING_OPTIONS:
        try:
            df = pd.read_csv(
                io.BytesIO(data),
                quoting=quoting,
                **base_kwargs
            )
            
            if not df.empty:
                LOGGER.debug(f"CSV parsed successfully with quoting={quoting}")
                return df
        except Exception as e:
            LOGGER.debug(f"CSV parse attempt failed (quoting={quoting}): {e}")
            continue
    
    # Próba 2: Python engine (wolniejszy, ale bardziej tolerancyjny)
    try:
        df = pd.read_csv(
            io.BytesIO(data),
            engine="python",
            **base_kwargs
        )
        
        if not df.empty:
            LOGGER.debug("CSV parsed successfully with python engine")
            return df
    except Exception as e:
        LOGGER.debug(f"Python engine failed: {e}")
    
    # Próba 3: Pozwól pandasowi zgadnąć wszystko
    try:
        df = pd.read_csv(
            io.BytesIO(data),
            encoding=encoding,
            low_memory=False,
            **_get_dtype_backend_kwargs(opts)
        )
        
        if not df.empty:
            LOGGER.warning("CSV parsed with fallback auto-detection")
            return df
    except Exception as e:
        LOGGER.error(f"All CSV parsing attempts failed: {e}")
        raise ValueError(f"Nie udało się sparsować CSV: {e}")
    
    raise ValueError("Nie udało się sparsować CSV po wszystkich próbach")


def _read_excel_robust(
    file_obj: io.BytesIO,
    opts: ParseOptions
) -> pd.DataFrame:
    """
    Robust Excel reader.
    
    Args:
        file_obj: BytesIO z danymi
        opts: Opcje parsowania
        
    Returns:
        DataFrame
        
    Raises:
        ValueError: Jeśli parsowanie się nie uda
    """
    backend_kwargs = _get_dtype_backend_kwargs(opts)
    
    # Przypadek 1: Konkretny arkusz wskazany
    if opts.excel_sheet is not None:
        try:
            df = pd.read_excel(
                file_obj,
                sheet_name=opts.excel_sheet,
                **backend_kwargs
            )
            LOGGER.debug(f"Excel sheet '{opts.excel_sheet}' loaded successfully")
            return df
        except Exception as e:
            raise ValueError(f"Nie udało się wczytać arkusza '{opts.excel_sheet}': {e}")
    
    # Przypadek 2: Wczytaj wszystkie arkusze (rzadko używane)
    if opts.excel_read_all_sheets:
        try:
            all_sheets = pd.read_excel(
                file_obj,
                sheet_name=None,
                **backend_kwargs
            )
            
            if isinstance(all_sheets, dict):
                # Połącz wszystkie arkusze
                dfs = [df for df in all_sheets.values() if not df.empty]
                
                if dfs:
                    combined = pd.concat(dfs, ignore_index=True)
                    LOGGER.info(f"Combined {len(dfs)} Excel sheets")
                    return combined
        except Exception as e:
            LOGGER.warning(f"Failed to read all sheets: {e}")
    
    # Przypadek 3: Automatyczny wybór największego arkusza
    try:
        all_sheets = pd.read_excel(
            file_obj,
            sheet_name=None,
            **backend_kwargs
        )
        
        if isinstance(all_sheets, dict):
            # Filtruj puste arkusze
            non_empty = [
                (name, df) for name, df in all_sheets.items()
                if isinstance(df, pd.DataFrame) and not df.empty
            ]
            
            if not non_empty:
                LOGGER.warning("All Excel sheets are empty")
                return pd.DataFrame()
            
            # Wybierz największy
            best_name, best_df = max(non_empty, key=lambda x: len(x[1]))
            LOGGER.info(f"Selected largest sheet: '{best_name}' ({len(best_df)} rows)")
            return best_df
        
        # Jeśli nie dict, zwróć bezpośrednio
        return all_sheets  # type: ignore
        
    except Exception as e:
        LOGGER.error(f"Excel parsing failed: {e}")
    
    # Fallback: Próba z openpyxl engine
    try:
        LOGGER.info("Trying openpyxl engine fallback")
        df = pd.read_excel(
            file_obj,
            engine="openpyxl",
            **backend_kwargs
        )
        return df
    except Exception as e:
        raise ValueError(f"Nie udało się wczytać pliku Excel: {e}")


def _read_json_robust(
    data: bytes,
    opts: ParseOptions
) -> pd.DataFrame:
    """
    Robust JSON reader z obsługą JSON Lines i zagnieżdżonych struktur.
    
    Args:
        data: Dane binarne
        opts: Opcje parsowania
        
    Returns:
        DataFrame
        
    Raises:
        ValueError: Jeśli parsowanie się nie uda
    """
    # Detekcja encoding
    encoding = _detect_encoding(data, opts.csv_try_encodings, opts.csv_sniff_bytes)
    
    try:
        text = data.decode(encoding, errors="replace").strip()
    except Exception:
        text = data.decode("utf-8", errors="replace").strip()
    
    if not text:
        raise ValueError("Pusty plik JSON")
    
    backend_kwargs = _get_dtype_backend_kwargs(opts)
    
    # Próba 1: JSON Lines (NDJSON)
    is_jsonlines = (
        opts.json_force_lines or
        (text.startswith("{") and "\n{" in text) or
        text.count("\n") > 0 and not text.startswith("[")
    )
    
    if is_jsonlines:
        try:
            df = pd.read_json(
                io.StringIO(text),
                lines=True,
                **backend_kwargs
            )
            LOGGER.debug("Parsed as JSON Lines")
            return df
        except Exception as e:
            LOGGER.debug(f"JSON Lines parsing failed: {e}")
    
    # Próba 2: Array of objects
    if text.startswith("[") and text.endswith("]"):
        try:
            data_obj = json.loads(text)
            
            if isinstance(data_obj, list) and data_obj:
                # Spróbuj normalizacji dla zagnieżdżonych struktur
                if opts.json_normalize_nested and isinstance(data_obj[0], dict):
                    try:
                        df = pd.json_normalize(
                            data_obj,
                            max_level=opts.json_max_level
                        )
                        LOGGER.debug("Parsed with json_normalize")
                        return df
                    except Exception:
                        pass
                
                # Fallback do DataFrame constructor
                df = pd.DataFrame(data_obj)
                LOGGER.debug("Parsed as array of objects")
                return df
        except Exception as e:
            LOGGER.debug(f"Array parsing failed: {e}")
    
    # Próba 3: Single object
    if text.startswith("{") and text.endswith("}"):
        try:
            data_obj = json.loads(text)
            
            if isinstance(data_obj, dict):
                df = pd.DataFrame([data_obj])
                LOGGER.debug("Parsed as single object")
                return df
        except Exception as e:
            LOGGER.debug(f"Single object parsing failed: {e}")
    
    # Próba 4: Fallback pandas read_json
    try:
        df = pd.read_json(
            io.StringIO(text),
            **backend_kwargs
        )
        LOGGER.warning("Parsed with fallback read_json")
        return df
    except Exception as e:
        raise ValueError(f"Nie udało się sparsować JSON: {e}")


def _read_parquet_robust(
    file_obj: io.BytesIO,
    opts: ParseOptions
) -> pd.DataFrame:
    """
    Robust Parquet reader.
    
    Args:
        file_obj: BytesIO z danymi
        opts: Opcje parsowania
        
    Returns:
        DataFrame
        
    Raises:
        ValueError: Jeśli parsowanie się nie uda
    """
    backend_kwargs = _get_dtype_backend_kwargs(opts)
    
    try:
        df = pd.read_parquet(file_obj, **backend_kwargs)
        LOGGER.debug("Parquet file loaded successfully")
        return df
    except Exception as e:
        raise ValueError(f"Nie udało się wczytać pliku Parquet: {e}")


def _read_feather_robust(
    file_obj: io.BytesIO,
    opts: ParseOptions
) -> pd.DataFrame:
    """
    Robust Feather reader.
    
    Args:
        file_obj: BytesIO z danymi
        opts: Opcje parsowania
        
    Returns:
        DataFrame
        
    Raises:
        ValueError: Jeśli parsowanie się nie uda
    """
    backend_kwargs = _get_dtype_backend_kwargs(opts)
    
    try:
        df = pd.read_feather(file_obj, **backend_kwargs)
        LOGGER.debug("Feather file loaded successfully")
        return df
    except Exception as e:
        raise ValueError(f"Nie udało się wczytać pliku Feather: {e}")


def _extract_pdf_text(file_obj: io.BytesIO, opts: ParseOptions) -> str:
    """
    Ekstrahuje tekst z PDF.
    
    Args:
        file_obj: BytesIO z danymi
        opts: Opcje parsowania
        
    Returns:
        Wyekstrahowany tekst
    """
    if not HAS_PDF or PdfReader is None:
        return "[Parser PDF niedostępny — zainstaluj PyPDF2: pip install PyPDF2]"
    
    try:
        reader = PdfReader(file_obj)
    except Exception as e:
        return f"[Błąd otwierania PDF: {e}]"
    
    if not reader.pages:
        return "[PDF nie zawiera stron]"
    
    parts: List[str] = []
    total_pages = len(reader.pages)
    
    LOGGER.info(f"Extracting text from {total_pages} PDF pages")
    
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text()
            
            if text:
                parts.append(text.strip())
            else:
                LOGGER.debug(f"Page {i+1} contains no extractable text")
        except Exception as e:
            LOGGER.warning(f"Failed to extract text from page {i+1}: {e}")
            parts.append(f"[Błąd ekstrakcji na stronie {i+1}]")
    
    result = "\n\n".join(parts).strip()
    
    if not result:
        return "[PDF nie zawiera tekstu do ekstrakcji]"
    
    LOGGER.info(f"Extracted {len(result)} characters from PDF")
    return result


def _extract_docx_text(file_obj: io.BytesIO, opts: ParseOptions) -> str:
    """
    Ekstrahuje tekst z DOCX.
    
    Args:
        file_obj: BytesIO z danymi
        opts: Opcje parsowania
        
    Returns:
        Wyekstrahowany tekst
    """
    if not HAS_DOCX or Document is None:
        return "[Parser DOCX niedostępny — zainstaluj python-docx: pip install python-docx]"
    
    try:
        doc = Document(file_obj)
    except Exception as e:
        return f"[Błąd otwierania DOCX: {e}]"
    
    if not doc.paragraphs:
        return "[DOCX nie zawiera paragrafów]"
    
    parts: List[str] = []
    
    for para in doc.paragraphs:
        if para and para.text:
            parts.append(para.text.strip())
    
    result = "\n".join(parts).strip()
    
    if not result:
        return "[DOCX nie zawiera tekstu]"
    
    LOGGER.info(f"Extracted {len(result)} characters from DOCX")
    return result


# ========================================================================================
# MAIN PARSING FUNCTIONS
# ========================================================================================

def parse_any_pro(
    file_name: str,
    file_bytes: bytes,
    opts: Optional[ParseOptions] = None
) -> ParseResult:
    """
    PRO API: Uniwersalny parser z pełnymi metadanymi.
    
    Args:
        file_name: Nazwa pliku (z rozszerzeniem)
        file_bytes: Dane binarne pliku
        opts: Opcje parsowania (opcjonalne)
        
    Returns:
        ParseResult z danymi i metadanymi
        
    Raises:
        ValueError: Dla błędów walidacji lub nieobsługiwanych formatów
    """
    # Walidacja podstawowa
    if not file_name:
        raise ValueError("Brak nazwy pliku")
    
    if file_bytes is None or len(file_bytes) == 0:
        raise ValueError("Pusty plik")
    
    # Domyślne opcje
    if opts is None:
        opts = ParseOptions()
    
    # Walidacja rozmiaru
    try:
        _validate_file_size(file_bytes, opts.max_file_size_mb, file_name)
    except ValueError as e:
        # Zwróć błąd jako ParseResult
        return ParseResult(
            df=None,
            text=None,
            warnings=[str(e)],
            meta={"error": str(e), "file_name": file_name}
        )
    
    # Inicjalizacja wyniku
    result = ParseResult(meta={
        "file_name": file_name,
        "size_bytes": len(file_bytes),
        "size_mb": round(len(file_bytes) / 1e6, 2)
    })
    
    # Detekcja rozszerzenia
    filename_lower = file_name.lower()
    
    # Obsługa kompresji
    if filename_lower.endswith(".csv.gz"):
        ext = "csv.gz"
    elif filename_lower.endswith(".tsv.gz"):
        ext = "tsv.gz"
    elif filename_lower.endswith(".json.gz"):
        ext = "json.gz"
    else:
        # Standardowe rozszerzenie
        parts = filename_lower.split(".")
        ext = parts[-1] if parts else ""
    
    # Walidacja rozszerzenia
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Nieobsługiwane rozszerzenie: .{ext}\n"
            f"Obsługiwane: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    
    result.meta["extension"] = ext
    
    LOGGER.info(f"Parsing file: {file_name} ({result.meta['size_mb']} MB, type: {ext})")
    
    try:
        # ============================================================================
        # COMPRESSED FILES
        # ============================================================================
        
        if ext in {"csv.gz", "tsv.gz", "json.gz"}:
            try:
                decompressed = gzip.decompress(file_bytes)
                LOGGER.debug(f"Decompressed: {len(file_bytes)} -> {len(decompressed)} bytes")
            except Exception as e:
                raise ValueError(f"Błąd dekompresji gzip: {e}")
            
            # Rekurencja z dekompresowanymi danymi
            base_ext = ext.replace(".gz", "")
            temp_filename = file_name.replace(".gz", "")
            
            return parse_any_pro(temp_filename, decompressed, opts)
        
        # ============================================================================
        # CSV / TSV / TXT
        # ============================================================================
        
        if ext in {"csv", "txt"}:
            df = _read_csv_robust(file_bytes, sep_hint=None, opts=opts)
            df = _validate_dataframe(df, opts)
            result.df = df
            result.meta["format"] = "csv"
            result.meta["shape"] = df.shape if df is not None else None
            
        elif ext == "tsv":
            df = _read_csv_robust(file_bytes, sep_hint="\t", opts=opts)
            df = _validate_dataframe(df, opts)
            result.df = df
            result.meta["format"] = "tsv"
            result.meta["shape"] = df.shape if df is not None else None
        
        # ============================================================================
        # EXCEL
        # ============================================================================
        
        elif ext in {"xlsx", "xls"}:
            buf = io.BytesIO(file_bytes)
            df = _read_excel_robust(buf, opts)
            df = _validate_dataframe(df, opts)
            result.df = df
            result.meta["format"] = "excel"
            result.meta["shape"] = df.shape if df is not None else None
        
        # ============================================================================
        # JSON / JSONL / NDJSON
        # ============================================================================
        
        elif ext in {"json", "jsonl", "ndjson"}:
            # NDJSON automatycznie włącza lines=True
            if ext in {"jsonl", "ndjson"}:
                opts = ParseOptions(**{**opts.__dict__, "json_force_lines": True})
            
            df = _read_json_robust(file_bytes, opts)
            df = _validate_dataframe(df, opts)
            result.df = df
            result.meta["format"] = "jsonl" if ext in {"jsonl", "ndjson"} else "json"
            result.meta["shape"] = df.shape if df is not None else None
        
        # ============================================================================
        # PARQUET
        # ============================================================================
        
        elif ext == "parquet":
            buf = io.BytesIO(file_bytes)
            df = _read_parquet_robust(buf, opts)
            df = _validate_dataframe(df, opts)
            result.df = df
            result.meta["format"] = "parquet"
            result.meta["shape"] = df.shape if df is not None else None
        
        # ============================================================================
        # FEATHER
        # ============================================================================
        
        elif ext == "feather":
            buf = io.BytesIO(file_bytes)
            df = _read_feather_robust(buf, opts)
            df = _validate_dataframe(df, opts)
            result.df = df
            result.meta["format"] = "feather"
            result.meta["shape"] = df.shape if df is not None else None
        
        # ============================================================================
        # PDF
        # ============================================================================
        
        elif ext == "pdf":
            buf = io.BytesIO(file_bytes)
            text = _extract_pdf_text(buf, opts)
            result.text = text
            result.meta["format"] = "pdf"
            result.meta["text_length"] = len(text) if text else 0
        
        # ============================================================================
        # DOCX
        # ============================================================================
        
        elif ext == "docx":
            buf = io.BytesIO(file_bytes)
            text = _extract_docx_text(buf, opts)
            result.text = text
            result.meta["format"] = "docx"
            result.meta["text_length"] = len(text) if text else 0
        
        # ============================================================================
        # LEGACY DOC (unsupported)
        # ============================================================================
        
        elif ext == "doc":
            result.add_warning(
                "Legacy .doc format jest przestarzały i nie jest w pełni obsługiwany. "
                "Proszę przekonwertować na .docx dla pełnej funkcjonalności."
            )
            result.text = (
                f"[Legacy .doc file detected: {len(file_bytes)} bytes]\n\n"
                "UWAGA: Format .doc nie jest w pełni obsługiwany.\n"
                "Proszę przekonwertować do .docx używając Microsoft Word lub LibreOffice."
            )
            result.meta["format"] = "doc"
            result.meta["legacy_format"] = True
        
        # ============================================================================
        # UNKNOWN (nie powinno się zdarzyć przez wcześniejszą walidację)
        # ============================================================================
        
        else:
            raise ValueError(f"Nieobsługiwany format: .{ext}")
        
    except ValueError:
        # Przepuść ValueError wyżej
        raise
    
    except Exception as e:
        # Wszystkie inne błędy
        error_msg = f"Błąd parsowania pliku {file_name}: {e}"
        LOGGER.error(error_msg, exc_info=True)
        
        result.add_warning(error_msg)
        result.meta["error"] = str(e)
        result.meta["error_type"] = type(e).__name__
        
        if opts.raise_on_warning:
            raise ValueError(error_msg) from e
    
    # Finalne logowanie
    if result.success:
        if result.is_tabular:
            LOGGER.info(
                f"✅ Parsed successfully: {result.meta.get('shape', 'N/A')} "
                f"({result.meta['format']})"
            )
        else:
            LOGGER.info(
                f"✅ Extracted text: {result.meta.get('text_length', 0)} chars "
                f"({result.meta['format']})"
            )
    else:
        LOGGER.warning(f"⚠️ Parsing completed with warnings: {file_name}")
    
    return result


def parse_any(
    file_name: str,
    file_bytes: bytes
) -> Tuple[Optional[pd.DataFrame], Optional[str]]:
    """
    Uniwersalny parser - BACKWARD COMPATIBLE API.
    
    Args:
        file_name: Nazwa pliku (z rozszerzeniem)
        file_bytes: Dane binarne pliku
        
    Returns:
        Tuple (DataFrame lub None, tekst lub None)
        
    Raises:
        ValueError: Dla błędów parsowania
        
    Examples:
        >>> df, text = parse_any("data.csv", csv_bytes)
        >>> df, text = parse_any("document.pdf", pdf_bytes)
    """
    result = parse_any_pro(file_name, file_bytes, ParseOptions())
    
    # Backward compatibility: zwróć tuple
    return (result.df, result.text)


# ========================================================================================
# HELPER FUNCTIONS FOR EXTERNAL USE
# ========================================================================================

def is_supported(file_name: str) -> bool:
    """
    Sprawdza czy plik jest obsługiwany.
    
    Args:
        file_name: Nazwa pliku
        
    Returns:
        True jeśli obsługiwany
        
    Examples:
        >>> is_supported("data.csv")
        True
        >>> is_supported("image.png")
        False
    """
    filename_lower = file_name.lower()
    
    # Obsługa kompresji
    if filename_lower.endswith(".csv.gz") or filename_lower.endswith(".tsv.gz") or filename_lower.endswith(".json.gz"):
        return True
    
    # Standardowe rozszerzenie
    ext = filename_lower.split(".")[-1] if "." in filename_lower else ""
    
    return ext in SUPPORTED_EXTENSIONS


def get_file_type(file_name: str) -> Optional[str]:
    """
    Zwraca typ pliku na podstawie rozszerzenia.
    
    Args:
        file_name: Nazwa pliku
        
    Returns:
        Typ pliku lub None jeśli nieobsługiwany
        
    Examples:
        >>> get_file_type("data.csv")
        'csv'
        >>> get_file_type("report.pdf")
        'pdf'
    """
    if not is_supported(file_name):
        return None
    
    filename_lower = file_name.lower()
    
    # Obsługa kompresji
    if filename_lower.endswith(".csv.gz"):
        return "csv"
    elif filename_lower.endswith(".tsv.gz"):
        return "tsv"
    elif filename_lower.endswith(".json.gz"):
        return "json"
    
    # Standardowe
    return filename_lower.split(".")[-1]


def detect_format_from_content(data: bytes, sample_size: int = 4096) -> Optional[str]:
    """
    Próbuje wykryć format pliku na podstawie zawartości (fallback gdy brak rozszerzenia).
    
    Args:
        data: Dane binarne
        sample_size: Rozmiar próbki do analizy
        
    Returns:
        Wykryty format lub None
        
    Examples:
        >>> detect_format_from_content(csv_bytes)
        'csv'
    """
    if len(data) == 0:
        return None
    
    sample = data[:sample_size]
    
    # PDF magic number
    if sample.startswith(b"%PDF"):
        return "pdf"
    
    # ZIP-based formats (XLSX, DOCX)
    if sample.startswith(b"PK\x03\x04"):
        # Bardziej szczegółowa detekcja
        if b"word/" in sample or b"[Content_Types].xml" in sample:
            return "docx"
        if b"xl/" in sample or b"worksheets/" in sample:
            return "xlsx"
        return "zip"
    
    # Parquet magic number
    if sample.startswith(b"PAR1"):
        return "parquet"
    
    # Feather/Arrow
    if b"ARROW" in sample[:100]:
        return "feather"
    
    # GZIP
    if sample.startswith(b"\x1f\x8b"):
        return "gzip"
    
    # JSON detection (simple heuristic)
    try:
        text = sample.decode("utf-8", errors="ignore").strip()
        
        if text.startswith("{") or text.startswith("["):
            # Spróbuj sparsować jako JSON
            try:
                json.loads(text[:1000])
                return "json"
            except Exception:
                pass
    except Exception:
        pass
    
    # CSV/TSV detection (text-based)
    try:
        text = sample.decode("utf-8", errors="ignore")
        
        # Sprawdź czy wygląda jak CSV
        if "," in text or ";" in text or "\t" in text:
            lines = text.split("\n")[:5]
            
            if len(lines) >= 2:
                # Sprawdź konsystencję separatorów
                for sep in [",", ";", "\t"]:
                    counts = [line.count(sep) for line in lines if line.strip()]
                    
                    if counts and len(set(counts)) == 1 and counts[0] > 0:
                        return "tsv" if sep == "\t" else "csv"
    except Exception:
        pass
    
    return None


def validate_file(file_name: str, file_bytes: bytes) -> Dict[str, Any]:
    """
    Waliduje plik bez parsowania całości.
    
    Args:
        file_name: Nazwa pliku
        file_bytes: Dane binarne
        
    Returns:
        Dict z informacjami o walidacji
        
    Examples:
        >>> info = validate_file("data.csv", csv_bytes)
        >>> print(info['valid'])
        True
    """
    info = {
        "file_name": file_name,
        "size_bytes": len(file_bytes),
        "size_mb": round(len(file_bytes) / 1e6, 2),
        "valid": False,
        "errors": [],
        "warnings": [],
        "detected_format": None,
        "supported": False
    }
    
    # Sprawdź czy pusty
    if len(file_bytes) == 0:
        info["errors"].append("Plik jest pusty")
        return info
    
    # Sprawdź rozszerzenie
    info["supported"] = is_supported(file_name)
    
    if not info["supported"]:
        info["errors"].append(f"Nieobsługiwane rozszerzenie")
        
        # Spróbuj wykryć format z zawartości
        detected = detect_format_from_content(file_bytes)
        
        if detected:
            info["detected_format"] = detected
            info["warnings"].append(
                f"Plik może być w formacie {detected}, ale rozszerzenie nie pasuje"
            )
    
    # Sprawdź rozmiar
    if info["size_mb"] > DEFAULT_MAX_FILE_SIZE_MB:
        info["warnings"].append(
            f"Plik przekracza zalecany rozmiar ({info['size_mb']} MB > {DEFAULT_MAX_FILE_SIZE_MB} MB)"
        )
    
    # Jeśli brak błędów - valid
    info["valid"] = len(info["errors"]) == 0
    
    return info


# ========================================================================================
# BATCH PROCESSING
# ========================================================================================

def parse_multiple(
    files: List[Tuple[str, bytes]],
    opts: Optional[ParseOptions] = None,
    combine_dataframes: bool = False
) -> List[ParseResult]:
    """
    Parsuje wiele plików naraz.
    
    Args:
        files: Lista tuple (nazwa_pliku, dane_binarne)
        opts: Opcje parsowania
        combine_dataframes: Czy połączyć wszystkie DataFrame'y w jeden
        
    Returns:
        Lista ParseResult dla każdego pliku
        
    Examples:
        >>> files = [("file1.csv", data1), ("file2.csv", data2)]
        >>> results = parse_multiple(files)
    """
    if opts is None:
        opts = ParseOptions()
    
    results: List[ParseResult] = []
    
    for file_name, file_bytes in files:
        try:
            result = parse_any_pro(file_name, file_bytes, opts)
            results.append(result)
        except Exception as e:
            LOGGER.error(f"Failed to parse {file_name}: {e}")
            
            # Dodaj error result
            error_result = ParseResult(
                df=None,
                text=None,
                warnings=[f"Błąd parsowania: {e}"],
                meta={
                    "file_name": file_name,
                    "error": str(e),
                    "error_type": type(e).__name__
                }
            )
            results.append(error_result)
    
    # Opcjonalne łączenie DataFrame'ów
    if combine_dataframes:
        dfs = [r.df for r in results if r.df is not None and not r.df.empty]
        
        if dfs:
            try:
                combined = pd.concat(dfs, ignore_index=True)
                LOGGER.info(f"Combined {len(dfs)} DataFrames into one ({combined.shape})")
                
                # Dodaj combined result
                combined_result = ParseResult(
                    df=combined,
                    text=None,
                    warnings=[],
                    meta={
                        "combined": True,
                        "source_files": len(dfs),
                        "shape": combined.shape
                    }
                )
                results.append(combined_result)
            except Exception as e:
                LOGGER.error(f"Failed to combine DataFrames: {e}")
    
    return results


# ========================================================================================
# EXPORT FUNCTIONS
# ========================================================================================

def get_capabilities() -> Dict[str, Any]:
    """
    Zwraca informacje o dostępnych funkcjach parsera.
    
    Returns:
        Dict z informacjami o capabilities
        
    Examples:
        >>> caps = get_capabilities()
        >>> print(caps['has_pdf'])
        True
    """
    return {
        "supported_extensions": sorted(SUPPORTED_EXTENSIONS),
        "has_chardet": HAS_CHARDET,
        "has_pdf": HAS_PDF,
        "has_docx": HAS_DOCX,
        "has_pyarrow": HAS_PYARROW,
        "pandas_version": pd.__version__,
        "max_file_size_mb": DEFAULT_MAX_FILE_SIZE_MB,
        "max_rows": DEFAULT_MAX_ROWS,
        "max_cols": DEFAULT_MAX_COLS
    }


def print_capabilities() -> None:
    """
    Wyświetla dostępne funkcje parsera.
    
    Examples:
        >>> print_capabilities()
        📦 Universal Parser Capabilities
        ...
    """
    caps = get_capabilities()
    
    print("📦 Universal Parser Capabilities\n")
    print(f"Pandas version: {caps['pandas_version']}")
    print(f"\nOptional dependencies:")
    print(f"  ✅ CharDet: {caps['has_chardet']}")
    print(f"  ✅ PyPDF2: {caps['has_pdf']}")
    print(f"  ✅ python-docx: {caps['has_docx']}")
    print(f"  ✅ PyArrow: {caps['has_pyarrow']}")
    print(f"\nSupported formats ({len(caps['supported_extensions'])}):")
    print(f"  {', '.join(caps['supported_extensions'])}")
    print(f"\nDefault limits:")
    print(f"  Max file size: {caps['max_file_size_mb']} MB")
    print(f"  Max rows: {caps['max_rows']:,}")
    print(f"  Max cols: {caps['max_cols']:,}")


# ========================================================================================
# TESTING UTILITIES
# ========================================================================================

def _create_test_csv() -> bytes:
    """Tworzy testowe dane CSV."""
    return b"name,age,city\nJohn,30,NYC\nJane,25,LA\n"


def _create_test_json() -> bytes:
    """Tworzy testowe dane JSON."""
    return b'[{"name":"John","age":30},{"name":"Jane","age":25}]'


def test_parser(verbose: bool = True) -> Dict[str, bool]:
    """
    Testuje parser z różnymi formatami.
    
    Args:
        verbose: Czy wyświetlać szczegóły
        
    Returns:
        Dict z wynikami testów
        
    Examples:
        >>> results = test_parser(verbose=False)
        >>> all(results.values())
        True
    """
    results = {}
    
    # Test CSV
    try:
        df, _ = parse_any("test.csv", _create_test_csv())
        results["csv"] = df is not None and len(df) == 2
        if verbose:
            print(f"✅ CSV: {results['csv']}")
    except Exception as e:
        results["csv"] = False
        if verbose:
            print(f"❌ CSV: {e}")
    
    # Test JSON
    try:
        df, _ = parse_any("test.json", _create_test_json())
        results["json"] = df is not None and len(df) == 2
        if verbose:
            print(f"✅ JSON: {results['json']}")
    except Exception as e:
        results["json"] = False
        if verbose:
            print(f"❌ JSON: {e}")
    
    # Test TSV
    try:
        tsv_data = b"name\tage\nJohn\t30\nJane\t25\n"
        df, _ = parse_any("test.tsv", tsv_data)
        results["tsv"] = df is not None and len(df) == 2
        if verbose:
            print(f"✅ TSV: {results['tsv']}")
    except Exception as e:
        results["tsv"] = False
        if verbose:
            print(f"❌ TSV: {e}")
    
    return results


# ========================================================================================
# MAIN
# ========================================================================================

if __name__ == "__main__":
    print_capabilities()
    print("\n" + "="*60 + "\n")
    print("Running tests...\n")
    test_results = test_parser(verbose=True)
    print(f"\n✅ Passed: {sum(test_results.values())}/{len(test_results)}")